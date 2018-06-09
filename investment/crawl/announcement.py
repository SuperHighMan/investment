#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# @File  : announcement.py
# @Author: Hui
# @Date  : 2018/6/5
# @Desc  : 此代码用于分析投资者关系活动记录表
#---------注：此代码只能在windows上运行

import requests
import docx
import os
import jieba
import jieba.analyse
from win32com import client as wc


from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LTTextBoxHorizontal, LAParams
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
from pdfminer.pdfparser import PDFParser, PDFDocument

PATH = 'C:/tmp'

def pdf_parse(_path):
    fp = open(_path, 'rb')  # rb以二进制读模式打开本地pdf文件
    #request = Request(url=_path, headers={'User-Agent': random.choice(user_agent)})  # 随机从user_agent列表中抽取一个元素
    #fp = urlopen(request) #打开在线PDF文档

    # 用文件对象来创建一个pdf文档分析器
    praser_pdf = PDFParser(fp)

    # 创建一个PDF文档
    doc = PDFDocument()

    # 连接分析器 与文档对象
    praser_pdf.set_document(doc)
    doc.set_parser(praser_pdf)

    # 提供初始化密码doc.initialize("123456")
    # 如果没有密码 就创建一个空的字符串
    doc.initialize()

    # 检测文档是否提供txt转换，不提供就忽略
    if not doc.is_extractable:
        raise PDFTextExtractionNotAllowed
    else:
        # 创建PDf资源管理器 来管理共享资源
        rsrcmgr = PDFResourceManager()

        # 创建一个PDF参数分析器
        laparams = LAParams()

        # 创建聚合器
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)

        # 创建一个PDF页面解释器对象
        interpreter = PDFPageInterpreter(rsrcmgr, device)

        # 循环遍历列表，每次处理一页的内容
        # doc.get_pages() 获取page列表
        for page in doc.get_pages():
            # 使用页面解释器来读取
            interpreter.process_page(page)

            # 使用聚合器获取内容
            layout = device.get_result()

            # 这里layout是一个LTPage对象 里面存放着 这个page解析出的各种对象 一般包括LTTextBox, LTFigure, LTImage, LTTextBoxHorizontal 等等 想要获取文本就获得对象的text属性，
            for out in layout:
                # 判断是否含有get_text()方法，图片之类的就没有
                # if hasattr(out,"get_text"):
                if isinstance(out, LTTextBoxHorizontal):

                    results = out.get_text()
                    print("results: " + results)

def load_key_word(file):
    """
    读取金融字典
    param: file: 文件
    return: list
    """
    with open(file, encoding='utf8') as dict:
        keys = [key.strip() for key in dict]

    return keys

def doc2docx(path, file):
    """
    MS doc文档转换成docx文档，注:此函数只能在Windows上运行
    return:
           str: 转换完成后的路径
    """
    word = wc.Dispatch('Word.Application')
    doc = word.Documents.Open(os.path.join(path, file))
    doc.SaveAs(os.path.join(path, 'tmp.docx'), 12, False, "", True, "", False, False, False, False)
    doc.Close()
    word.Quit()
    return os.path.join(path, 'tmp.docx')

def seperate_company(doc, dict):
    """
    分离文档中的涉案人员 一般对应投资者关系活动记录表中的“参与单位名称及人员姓名”一栏
    param:
    return:
           int : 返回参与单位的数量
    """
    jieba.load_userdict(dict)
    keys = load_key_word(dict)
    # 处理MS文档中的表格
    tables = [table for table in doc.tables];
    """
    ii = 0
    for table in tables:
        for row in table.rows:
            print(ii)
            for cell in row.cells:
                print(cell.text)
                print('haha')
            ii = ii+1
        print('\n')
    """
    # 按照文档读取参加人员的信息
    if len(tables) > 0:
        sum = 0
        row = tables[0].rows[1]
        for cell in row.cells:
            #print(cell.text)
            words = jieba.cut(cell.text)
            #print(','.join(words))
            for word in words:
                if word in keys:
                    sum = sum + 1
        #print(u'此份报告涉及投资公司:%d家'%sum)
        return sum
    else:
        return 0

def get_announcement_from_juchao(code, pause=0.02):
    """
    获取巨潮资讯网站上投资者活动关系
    :param code: string e.g. 000860
    :param pause:
    :return:
    """
    url = "http://www.cninfo.com.cn/cninfo-new/announcement/query"
    data = {'stock':code,
            'searchkey':'',
            'plate':'',
            'category':'',
            'trade':'',
            'column':'szse',
            'columnTitle':u'历史公告查询',
            'pageNum':1,
            'pageSize':30,
            'tabName':'relation',
            'sortName':'',
            'sortType':'',
            'limit':'',
            'showTitle':u'',
            'seDate':u'请选择日期'}
    response = requests.post(url, data)
    totalAnnouncement = response.json()['totalAnnouncement']
    content = response.json()['announcements']
    for element in content:
        title = element['announcementTitle']
        adjunctUrl = element['adjunctUrl']
        adjunctType = element['adjunctType']
        #print(u'主题:%s,URL:%s,文件类型:%s' % (title, adjunctUrl, adjunctType))
        file_url = "http://www.cninfo.com.cn/" + adjunctUrl
        r = requests.get(file_url)
        with open(os.path.join(PATH, 'tmp.'+adjunctType), 'wb') as file:
            file.write(r.content)
        # 读取MS文档
        if adjunctType == 'DOC':
            doc = docx.Document(doc2docx(PATH, 'tmp.'+adjunctType))
            totalSum = seperate_company(doc, 'dict.txt')
            yield ('%s,%s,%s,%d\n'%(code, adjunctUrl.split('/')[1], adjunctType, totalSum))
        elif adjunctType == 'DOCX':
            doc = docx.Document(os.path.join(PATH, 'tmp.'+adjunctType))
            totalSum = seperate_company(doc, 'dict.txt')
            yield ('%s,%s,%s,%d\n'%(code, adjunctUrl.split('/')[1], adjunctType, totalSum))
        #content = '\n'.join([para.text for para in doc.paragraphs])
        #print(content)
    #print(totalAnnouncement)

if __name__ == '__main__':
    codes = ['002507','002415','002236','000725','600183','300136','002241','002475','600703','600309','600352']
    with open(os.path.join(PATH, 'crawl.txt'), 'wb') as out:
        for code in codes:
            generator = get_announcement_from_juchao(code)
            for i in generator:
                print(i)
                out.write(bytes(i, encoding='utf8'))
    print('crawl finish')
    #pdf_parse('C:/Users/Hui/Desktop/1205035254.PDF')
